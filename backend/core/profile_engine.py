"""
Profile Intelligence Engine.

Handles:
- Resume text extraction (PDF/DOCX)
- LLM-powered structured profile parsing
- Vulnerability analysis (strengths, weaknesses, pressure points)
"""

import logging
from pathlib import Path

from core.llm_provider import llm
from models.profile import ParsedProfile, ProfileIntelligence

logger = logging.getLogger(__name__)


# ──────────────────────────────────────────────────────────────
# Candidate Type Classification (Deterministic — NOT LLM-based)
# ──────────────────────────────────────────────────────────────

# Titles that are NEVER professional experience regardless of duration
_INTERN_KEYWORDS = [
    "intern", "internship", "summer intern", "winter intern",
    "trainee intern", "student intern", "research intern",
]

_STUDENT_KEYWORDS = [
    "campus ambassador", "volunteer", "nss", "ncc",
    "club president", "club member", "student coordinator",
    "event organizer", "college representative", "hackathon",
    "teaching assistant", "student researcher", "student",
]

# Titles that ARE professional experience
_PROFESSIONAL_KEYWORDS = [
    "engineer", "developer", "analyst", "consultant", "associate",
    "executive", "manager", "lead", "architect", "designer",
    "accountant", "scientist", "founder", "co-founder",
    "entrepreneur", "freelance", "contract", "officer",
    "coordinator", "specialist", "director", "head",
    "vice president", "vp", "ceo", "cto", "cfo", "coo",
]

# Trainee programs that count as professional
_PROFESSIONAL_TRAINEE = [
    "management trainee", "graduate trainee",
    "graduate engineer trainee", "get",
]


def _parse_duration_months(duration_str: str) -> float:
    """Best-effort parse of a duration string into months."""
    import re
    duration_str = duration_str.lower().strip()

    # Try "X years Y months" pattern
    years = 0.0
    months = 0.0
    y_match = re.search(r"(\d+\.?\d*)\s*(?:year|yr|y)", duration_str)
    m_match = re.search(r"(\d+\.?\d*)\s*(?:month|mo|m\b)", duration_str)
    if y_match:
        years = float(y_match.group(1))
    if m_match:
        months = float(m_match.group(1))
    if y_match or m_match:
        return years * 12 + months

    # Try "X weeks" pattern
    w_match = re.search(r"(\d+)\s*(?:week|wk|w)", duration_str)
    if w_match:
        return float(w_match.group(1)) / 4.0

    # Try date range "Jan 2022 - Jun 2023"
    date_pattern = r"(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)\w*\.?\s*\d{4}"
    dates = re.findall(date_pattern, duration_str)
    if len(dates) >= 2:
        from datetime import datetime
        try:
            start = datetime.strptime(re.sub(r"(\w{3})\w*\.?", r"\1", dates[0]), "%b %Y")
            end = datetime.strptime(re.sub(r"(\w{3})\w*\.?", r"\1", dates[1]), "%b %Y")
            diff = (end.year - start.year) * 12 + (end.month - start.month)
            return max(diff, 0)
        except Exception:
            pass

    return 0.0


def _is_intern_title(role: str) -> bool:
    """Check if a role title is an internship."""
    role_lower = role.lower().strip()
    # Direct keyword match
    for kw in _INTERN_KEYWORDS:
        if kw in role_lower:
            return True
    return False


def _is_student_activity(role: str) -> bool:
    """Check if a role is a student/volunteer activity."""
    role_lower = role.lower().strip()
    for kw in _STUDENT_KEYWORDS:
        if kw in role_lower:
            return True
    return False


def _is_professional_trainee(role: str) -> bool:
    """Check if a trainee role counts as professional."""
    role_lower = role.lower().strip()
    for kw in _PROFESSIONAL_TRAINEE:
        if kw in role_lower:
            return True
    return False


def classify_candidate_type(parsed_profile: dict) -> dict:
    """
    Deterministic classification of candidate as Fresher or Experienced Professional.

    Returns dict with:
        - type: "fresher" or "experienced"
        - confidence: "high", "medium", or "low"
        - reason: explanation string
        - professional_months: total months of professional experience
    """
    work_experience = parsed_profile.get("work_experience", [])
    internships = parsed_profile.get("internships", [])

    professional_months = 0.0
    professional_roles = []
    intern_roles = []
    ambiguous_roles = []

    # Step 1: Classify all work_experience entries
    for entry in work_experience:
        if not isinstance(entry, dict):
            continue
        role = str(entry.get("role", "")).strip()
        company = str(entry.get("company", "")).strip()
        duration = str(entry.get("duration", "")).strip()
        responsibilities = entry.get("responsibilities", [])
        achievements = entry.get("achievements", [])
        months = _parse_duration_months(duration)

        # Check all text fields for intern signals
        all_text = f"{role} {company} {' '.join(responsibilities)} {' '.join(achievements)}".lower()
        has_intern_signal = any(kw in all_text for kw in _INTERN_KEYWORDS)

        if _is_intern_title(role) or has_intern_signal:
            # Internship incorrectly placed in work_experience
            intern_roles.append(role)
        elif _is_student_activity(role):
            # Student activity, skip
            continue
        elif _is_professional_trainee(role):
            # Management/Graduate trainee — counts
            professional_months += months if months > 0 else 6
            professional_roles.append(role)
        else:
            # Check if it has a recognizable professional title
            role_lower = role.lower()
            is_prof = any(kw in role_lower for kw in _PROFESSIONAL_KEYWORDS)

            if is_prof and months >= 6:
                # Clear professional role with substantial duration
                professional_months += months
                professional_roles.append(role)
            elif is_prof and months > 0 and months < 6:
                # Short-duration role with professional title — likely internship
                # (e.g., "Analyst" for 2 months is probably an internship)
                intern_roles.append(role)
            elif is_prof and months == 0:
                # Professional title but no duration info — count cautiously
                ambiguous_roles.append(role)
                professional_months += 3  # conservative
            elif months > 0 and months <= 6:
                # Short duration + no professional title = almost certainly internship
                intern_roles.append(role)
            elif months > 6:
                # Long duration but unrecognized title — probably professional
                ambiguous_roles.append(role)
                professional_months += months
            else:
                # No title match, no duration — default to intern (fresher-first)
                intern_roles.append(role)

    # Step 2: Internships array — NEVER counts
    for entry in internships:
        if isinstance(entry, dict):
            intern_roles.append(str(entry.get("role", "Intern")))

    # Step 3: Decision
    if professional_months >= 6 and len(professional_roles) > 0:
        confidence = "high" if professional_months >= 12 else "medium"
        return {
            "type": "experienced",
            "confidence": confidence,
            "reason": f"{len(professional_roles)} professional role(s) totaling ~{professional_months:.0f} months",
            "professional_months": professional_months,
        }
    elif len(ambiguous_roles) > 0 and professional_months >= 6:
        return {
            "type": "experienced",
            "confidence": "low",
            "reason": f"Ambiguous roles ({', '.join(ambiguous_roles)}) with ~{professional_months:.0f} months",
            "professional_months": professional_months,
        }
    else:
        reason_parts = []
        if intern_roles:
            reason_parts.append(f"{len(intern_roles)} internship(s)")
        if not work_experience and not internships:
            reason_parts.append("no work history")
        elif not professional_roles:
            reason_parts.append("no professional roles found")
        reason = "Fresher: " + ", ".join(reason_parts) if reason_parts else "Fresher: no professional experience"
        return {
            "type": "fresher",
            "confidence": "high",
            "reason": reason,
            "professional_months": professional_months,
        }


async def extract_text_from_file(file_path: str) -> str:
    """Extract text from PDF or DOCX files. Uses OCR fallback for image-only PDFs."""
    path = Path(file_path)
    ext = path.suffix.lower()
    text = ""

    if ext == ".pdf":
        text = _extract_pdf(file_path)
        if not text.strip():
            logger.info(f"No text extracted from {file_path}, attempting OCR fallback...")
            text = await _extract_pdf_ocr_fallback(file_path)
    elif ext == ".docx":
        text = _extract_docx(file_path)
    elif ext == ".doc":
        raise ValueError("Legacy .doc format is not supported. Please convert to .docx or .pdf.")
    elif ext == ".txt":
        text = path.read_text(encoding="utf-8")
    else:
        raise ValueError(f"Unsupported file format: {ext}")
        
    return text

async def _extract_pdf_ocr_fallback(file_path: str) -> str:
    """Fallback OCR using free ocr.space API for image-only PDFs."""
    import httpx
    try:
        async with httpx.AsyncClient(timeout=60) as client:
            with open(file_path, 'rb') as f:
                files = {'file': (Path(file_path).name, f, 'application/pdf')}
                data = {'apikey': 'helloworld', 'language': 'eng', 'isOverlayRequired': 'false'}
                response = await client.post('https://api.ocr.space/parse/image', data=data, files=files)
                
                if response.status_code == 200:
                    result = response.json()
                    parsed_results = result.get("ParsedResults", [])
                    if parsed_results:
                        texts = [res.get("ParsedText", "") for res in parsed_results]
                        return "\n".join(texts).strip()
    except Exception as e:
        logger.warning(f"OCR fallback failed: {e}")
    return ""


def _extract_pdf(file_path: str) -> str:
    """Extract text from PDF using PyMuPDF."""
    import fitz  # PyMuPDF

    text_parts = []
    with fitz.open(file_path) as doc:
        for page in doc:
            text_parts.append(page.get_text())

    return "\n".join(text_parts).strip()


def _extract_docx(file_path: str) -> str:
    """Extract text from DOCX including tables. Falls back to raw read if python-docx not available."""
    try:
        import docx
        doc = docx.Document(file_path)
        text_parts = []
        for p in doc.paragraphs:
            if p.text.strip():
                text_parts.append(p.text.strip())
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text.strip())
        return "\n".join(text_parts).strip()
    except ImportError:
        logger.warning("python-docx not installed. Attempting raw text extraction.")
        return Path(file_path).read_text(encoding="utf-8", errors="ignore")


async def parse_profile(raw_text: str) -> dict:
    """
    Use LLM to extract structured profile from resume text.
    Returns a ParsedProfile dict.
    """
    system_prompt = """You are a precise resume parsing engine for IIM MBA admission interviews.

Extract ALL relevant information from the candidate's resume into the structured JSON format.

EXTRACTION RULES:
- Extract EVERY education entry — 10th, 12th, graduation, post-graduation. Include degree, field of study, institution name, year of passing, and score (CGPA or percentage). Do NOT leave institution or field empty if it's mentioned anywhere in the resume.
- Extract full-time work experience ONLY under "work_experience". Include company, role, duration, responsibilities, and achievements.
- Extract internships SEPARATELY under "internships" — NOT under work_experience. Internships are short-term (typically < 6 months), student-level work placements. Do not confuse them with full-time jobs.
- Extract skills, certifications, hobbies, extracurriculars, and achievements
- If CAT score is mentioned, extract it. If NOT mentioned, leave it as empty string — do NOT flag its absence.
- If career goals or MBA motivation is stated, extract it. If NOT mentioned, leave as empty string.
- If information is not present in the resume, leave the field empty or as an empty list
- Do NOT fabricate or infer information that isn't in the resume
- Read the ENTIRE resume carefully. Many resumes list education at the bottom — make sure you don't miss it.
"""

    result = await llm.generate(
        system_prompt=system_prompt,
        user_message=f"Parse this resume:\n\n{raw_text}",
        schema=ParsedProfile,
        temperature=0.1,  # Low temperature for precise extraction
    )

    return result if isinstance(result, dict) else result.model_dump() if hasattr(result, 'model_dump') else result


async def analyze_profile(parsed_profile: dict) -> dict:
    """
    Use LLM to generate intelligence report: strengths, weaknesses,
    pressure points, and likely panel questions.
    """
    import json

    system_prompt = """You are a senior IIM interview panel preparation system.

Given a candidate's parsed profile, generate a comprehensive intelligence report.

CRITICAL RULES — READ BEFORE ANALYZING:
1. FRESHER vs EXPERIENCED: If "work_experience" is EMPTY or contains ZERO entries, the candidate is a FRESHER. Internships do NOT count as work experience. Do NOT flag "lack of work experience" as a weakness for a fresher — that's expected.
2. DO NOT flag the absence of CAT score, target IIM list, or admission test scores as a weakness. These are NEVER included in resumes. That's not how resumes work.
3. DO NOT flag the absence of career goals unless the candidate explicitly wrote vague or contradictory goals. If the field is simply empty, ignore it — resumes often don't include career goal statements.
4. Every pressure point MUST reference something the candidate ACTUALLY wrote. Do not invent hypothetical gaps. If the resume says they interned somewhere, probe that internship. If their CGPA is low, probe that. But don't say "unexplained future work entry" when it's clearly an internship.
5. INTERNSHIPS listed separately from work_experience are clearly labeled as internships. Treat them as student-level learning experiences, NOT as unexplained job entries.

WHAT TO ANALYZE:
1. STRENGTHS: What the panel might acknowledge or find impressive based on ACTUAL resume content. **High scores (≥80% or ≥8.0 CGPA) are strengths, NOT pressure points.**
2. WEAKNESSES: Real gaps or inconsistencies PRESENT in the data (low scores, skill mismatches, career gaps between dated entries)
3. PRESSURE POINTS: Specific areas where the panel can create maximum pressure based on what the candidate ACTUALLY claims. **These must be actionable interview questions**, not abstract commentary. Each pressure point should be phrased as a specific question the panel would ask.
4. LIKELY QUESTIONS: 10-15 specific questions the panel would likely ask THIS candidate based on their actual background
5. CAREER TRANSITION RISKS: ONLY if the candidate has work experience in one field and is switching to another
6. TECHNICAL DEEP DIVES: Specific technical areas to probe based on their stated skills and academic background

EXAMPLES OF GOOD PRESSURE POINTS (based on actual data):
- "With a CGPA of 6.8, how do you justify this lower score and reassure the panel about your academic rigor?"
- "Regarding your internship at [Company], what were the specific measurable outcomes of your work, and how did you add tangible business value?"
- "You claim leadership in [Activity] — can you provide specific examples of difficult decisions you made and their impact?"

EXAMPLES OF BAD PRESSURE POINTS (DO NOT generate these):
- "Explaining the BBA 92% score – why it is not 100%" ← High scores are strengths, not pressure points
- "Absence of CAT score" ← Never on resumes
- "No target IIM list" ← Never on resumes
- "Unexplained future work entry" ← That's an internship, not unexplained
- "Lack of full-time work experience" ← Normal for a fresher

Be ruthless but FAIR. Only probe what's actually there.
"""

    result = await llm.generate(
        system_prompt=system_prompt,
        user_message=f"Analyze this candidate profile:\n\n{json.dumps(parsed_profile, indent=2)}",
        schema=ProfileIntelligence,
        temperature=0.3,
    )

    data = result if isinstance(result, dict) else result.model_dump() if hasattr(result, 'model_dump') else result

    # Robust flattening in case the LLM returned the schema structure itself
    flat = {}
    source = data.get("properties", data) if isinstance(data, dict) else {}
    for key in ["strengths", "weaknesses", "pressure_points", "likely_questions", "career_transition_risks", "technical_deep_dives"]:
        val = source.get(key, [])
        if isinstance(val, dict):
            if "items" in val:
                flat[key] = val["items"]
            elif "default" in val:
                flat[key] = val["default"]
            else:
                flat[key] = []
        elif isinstance(val, list):
            flat[key] = val
        else:
            flat[key] = []
            
    return flat
